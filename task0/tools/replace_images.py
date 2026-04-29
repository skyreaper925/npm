# -*- coding: utf-8 -*-
"""Автоматическое обновление отчёта

Что делает:
  1. Подменяет 8 изображений (word/media/image1..8.png) на актуальные
     версии из data/img/.
  2. Добавляет/обновляет в конце документа автоматический блок «Расчётные
     значения» с параметрами в начальный и последний момент времени из
     data/{tag}.out и процентом пропусков из data/{tag}.skip.

Текстовые правки по комментариям преподавателя НЕ выполняются автоматически
(см. docs/правки.md для готовых формулировок — вставлять руками в Word).

Использование (из любой директории):
    python tools/replace_images.py
"""
import io
import os
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

import numpy as np

# В Windows консоли по умолчанию cp1251 — не умеет кодировать русский в выводе.
# Принудительно делаем stdout UTF-8 (стандартный приём для Python 3).
if hasattr(sys.stdout, "buffer") and sys.stdout.encoding.lower() not in ("utf-8", "utf8"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# === Пути от корня task0/ =================================================
TASK_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = TASK_ROOT / "docs"
IMG_DIR = TASK_ROOT / "data" / "img"  # notebook всегда рендерит сюда

# Где искать .out / .skip от relax.py: SLURM на кластере не позволяет
# создавать подпапки, поэтому файлы могут оказаться в src/, а не в data/.
# Скрипт пробует кандидатов по очереди и берёт первого, у кого есть 05.out.
DATA_CANDIDATES = [
    TASK_ROOT / "data",
    TASK_ROOT / "src",
    TASK_ROOT,
]

SRC = DOCS_DIR / "Володин_0.docx"
DST = DOCS_DIR / "Володин_0.docx"


def find_data_dir(override=None):
    """Возвращает первую папку из кандидатов, где есть 05.out."""
    candidates = [Path(override)] if override else DATA_CANDIDATES
    for c in candidates:
        if (c / "05.out").exists():
            return c
    # ничего не нашли — вернём первый кандидат (там скрипт ничего не запишет
    # и просто покажет "Расчётные данные не найдены").
    return candidates[0]


# === Соответствие image{N}.png в docx → файл из data/img/ =================
# Если порядок в твоём отчёте другой — отредактируй словарь ниже.
IMAGE_MAP = {
    "word/media/image1.png": "Txx_combined.png",
    "word/media/image2.png": "H_combined.png",
    "word/media/image3.png": "T_combined.png",
    "word/media/image4.png": "anisotropy_log.png",
    "word/media/image5.png": "f_marginal_evolution.png",
    "word/media/image6.png": "f_xy_u05.png",
    "word/media/image7.png": "f_xy_u10.png",
    "word/media/image8.png": "f_xy_u15.png",
}

# === Маркер автоблока в конце документа ===================================
AUTO_BLOCK_MARKER = "[АВТО-ВСТАВКА: расчётные значения]"


def load_skip_data(data_dir):
    skip = {}
    for u in (0.5, 1.0, 1.5):
        tag = f"{int(round(u * 10)):02d}"
        path = data_dir / f"{tag}.skip"
        if not path.exists():
            continue
        kv = {}
        with open(path, encoding="utf-8") as f:
            for line in f:
                if "=" in line:
                    k, v = line.strip().split("=", 1)
                    kv[k] = v
        skip[u] = {
            "percent": float(kv.get("skip_percent", 0)),
            "actual_steps": int(kv.get("actual_steps", 0)),
            "elapsed_seconds": float(kv.get("elapsed_seconds", 0)) or None,
        }
    return skip


def load_out_summary(data_dir):
    s = {}
    for u in (0.5, 1.0, 1.5):
        tag = f"{int(round(u * 10)):02d}"
        path = data_dir / f"{tag}.out"
        if not path.exists():
            continue
        arr = np.loadtxt(path)
        s[u] = {
            "rows": arr.shape[0],
            "T_start": arr[0, 0], "T_xx_start": arr[0, 1], "H_start": arr[0, 2],
            "T_end": arr[-1, 0], "T_xx_end": arr[-1, 1], "H_end": arr[-1, 2],
        }
    return s


def replace_images(temp_dir):
    for inner, src_name in IMAGE_MAP.items():
        full_inner = Path(temp_dir) / Path(inner)
        full_src = IMG_DIR / src_name
        if not full_src.exists():
            print(f"  [SKIP] нет PNG: {full_src}")
            continue
        if not full_inner.exists():
            print(f"  [SKIP] нет якоря: {inner}")
            continue
        shutil.copy(full_src, full_inner)
        print(f"  [IMG]  {inner} ← {src_name}")


def add_auto_block(dst_path, summary, skip_data):
    """Добавляет/обновляет в конце документа таблицу с расчётными значениями.
    Идемпотентен: ищет AUTO_BLOCK_MARKER, если есть — удаляет всё от него
    и до конца, затем добавляет новый блок."""
    from docx import Document

    doc = Document(str(dst_path))
    body = doc.element.body
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    # Идём по прямым детям body (sectPr должен остаться последним) и ищем
    # параграф с маркером. После него всё удаляем — это наш авто-блок.
    direct_children = list(body)
    cut_idx = None
    for i, child in enumerate(direct_children):
        if child.tag == f"{ns}p":
            text = "".join((t.text or "") for t in child.iter(f"{ns}t"))
            if AUTO_BLOCK_MARKER in text:
                cut_idx = i
                break
    if cut_idx is not None:
        # sectPr (свойства секции, типа размера страницы) — последний дочерний;
        # его не удаляем.
        removed = 0
        for child in direct_children[cut_idx:]:
            if child.tag == f"{ns}sectPr":
                continue
            body.remove(child)
            removed += 1
        print(f"  [AUTO] старый авто-блок удалён ({removed} элементов)")

    # Открывающий маркер
    h_marker = doc.add_paragraph()
    h_marker.add_run(AUTO_BLOCK_MARKER).italic = True

    if not summary:
        doc.add_paragraph(
            "Расчётные данные не найдены (нужно сначала запустить relax.py "
            "и поместить .out / .skip в data/)."
        )
    else:
        intro = doc.add_paragraph()
        intro.add_run(
            "Сводка по релаксации для трёх значений u (значения T, T_xx, H в начале и конце):"
        ).bold = True

        table = doc.add_table(rows=1, cols=9)
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        hdr = table.rows[0].cells
        for i, txt in enumerate([
            "u", "шагов",
            "T_xx (начало)", "T_xx (конец)",
            "T (const)",
            "H (начало)", "H (конец)",
            "пропуски, %",
            "время счёта",
        ]):
            hdr[i].text = txt

        for u in sorted(summary):
            s = summary[u]
            row = table.add_row().cells
            steps = (skip_data.get(u, {}).get("actual_steps")
                     if skip_data else s["rows"]) or s["rows"]
            pct = (skip_data.get(u, {}).get("percent")
                   if skip_data else None)
            elapsed = (skip_data.get(u, {}).get("elapsed_seconds")
                       if skip_data else None)
            row[0].text = f"{u}"
            row[1].text = f"{steps}"
            row[2].text = f"{s['T_xx_start']:.5f}"
            row[3].text = f"{s['T_xx_end']:.5f}"
            row[4].text = f"{s['T_end']:.5f}"
            row[5].text = f"{s['H_start']:.5f}"
            row[6].text = f"{s['H_end']:.5f}"
            row[7].text = f"{pct:.2f}" if pct is not None else "—"
            if elapsed:
                m, s_ = divmod(int(elapsed), 60)
                h, m = divmod(m, 60)
                # row[8].text = f"{h}ч {m}м {s_}с" if h else f"{m}м {s_}с"
                row[8].text = f"{h}ч {m}м" if h else f"{m}м"
            else:
                row[8].text = "—"

    # Закрывающий маркер
    doc.add_paragraph(AUTO_BLOCK_MARKER).runs[0].italic = True

    doc.save(str(dst_path))
    print(f"  [AUTO] новый блок добавлен")


def main():
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument("--data-dir", default=None,
                   help="папка с .out / .npz / .skip (по умолчанию ищет в data/, src/, корне)")
    args = p.parse_args()

    if not SRC.exists():
        sys.exit(f"Не найден исходный файл: {SRC}")

    data_dir = find_data_dir(args.data_dir)
    skip_data = load_skip_data(data_dir)
    summary = load_out_summary(data_dir)

    print(f"==> Источник: {SRC}")
    print(f"==> Цель:     {DST}")
    print(f"==> Данные:   {data_dir}")
    print(f"==> Найдены .out для u: {sorted(summary)}")
    print(f"==> Найдены .skip для u: {sorted(skip_data)}")
    print()

    # Шаг 1: распаковка SRC, замена картинок, перепаковка в DST
    with tempfile.TemporaryDirectory() as td:
        with zipfile.ZipFile(SRC, "r") as zin:
            zin.extractall(td)
        replace_images(td)
        with zipfile.ZipFile(DST, "w", zipfile.ZIP_DEFLATED) as zout:
            for root, _, files in os.walk(td):
                for fn in files:
                    full = os.path.join(root, fn)
                    rel = os.path.relpath(full, td).replace("\\", "/")
                    zout.write(full, rel)

    # Шаг 2: добавить авто-блок с числами (через python-docx)
    add_auto_block(DST, summary, skip_data)

    print()
    print(f"готово: {DST}")


if __name__ == "__main__":
    main()

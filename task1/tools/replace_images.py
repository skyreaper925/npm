# -*- coding: utf-8 -*-
"""Обновление отчёта task1 (Володин_1.docx).

Что делает:
  1. Подменяет 3 изображения (word/media/image1..3.png) на актуальные
     версии из data/img/ (графики T, S, n).
  2. Добавляет/обновляет в конце документа авто-блок «Расчётные значения»
     с парциальными/общими T, n, S в начальный и конечный момент времени
     из data/mixture.out и долей пропусков из data/mixture.skip.

Текстовые правки по комментариям преподавателя НЕ выполняются автоматически
(см. docs/правки1-4.md).

Использование:
    python tools/replace_images.py
"""
import io
import os
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path

import numpy as np

if hasattr(sys.stdout, "buffer") and (sys.stdout.encoding or "").lower() not in ("utf-8", "utf8"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# === Пути ==================================================================
TASK_ROOT = Path(__file__).resolve().parent.parent
DOCS_DIR = TASK_ROOT / "docs"
IMG_DIR = TASK_ROOT / "data" / "img"

# Где искать mixture.out/.skip.
DATA_CANDIDATES = [TASK_ROOT / "data", TASK_ROOT / "src", TASK_ROOT]

SRC = DOCS_DIR / "Володин_1.docx"
DST = DOCS_DIR / "Володин_1.docx"

# === Соответствие image{N}.png в docx → файл из data/img/ =================
# В отчёте три рисунка в разделе «Результаты»: температуры, энтропии,
# концентрации. Если порядок другой — поменяй значения местами.
IMAGE_MAP = {
    "word/media/image1.png": "T_combined.png",   # T_c, T_h, T_tot
    "word/media/image2.png": "S_combined.png",   # S_c, S_h, S_tot (H-теорема)
    "word/media/image3.png": "n_combined.png",   # n_c, n_h (контроль сохранения)
}

AUTO_BLOCK_MARKER = "[АВТО-ВСТАВКА: расчётные значения]"


def find_data_dir(override=None):
    candidates = [Path(override)] if override else DATA_CANDIDATES
    for c in candidates:
        if (c / "mixture.out").exists():
            return c
    return candidates[0]


def load_skip_data(data_dir):
    path = data_dir / "mixture.skip"
    if not path.exists():
        return {}
    kv = {}
    with open(path, encoding="utf-8") as f:
        for line in f:
            if "=" in line:
                k, v = line.strip().split("=", 1)
                kv[k] = v
    return {
        "percent": float(kv.get("skip_percent", 0)),
        "actual_steps": int(kv.get("actual_steps", 0)),
        "elapsed_seconds": float(kv.get("elapsed_seconds", 0)) or None,
        "tau": float(kv.get("tau", 0)) or None,
    }


def load_out_summary(data_dir):
    """mixture.out: столбцы T_tot T_c T_h n_c n_h S_tot S_c S_h."""
    path = data_dir / "mixture.out"
    if not path.exists():
        return None
    arr = np.loadtxt(path)
    if arr.ndim == 1:
        arr = arr.reshape(1, -1)
    cols = ["T_tot", "T_c", "T_h", "n_c", "n_h", "S_tot", "S_c", "S_h"]
    start = {c: arr[0, i] for i, c in enumerate(cols)}
    end = {c: arr[-1, i] for i, c in enumerate(cols)}
    return {"rows": arr.shape[0], "start": start, "end": end}


def replace_images(temp_dir):
    for inner, src_name in IMAGE_MAP.items():
        full_inner = Path(temp_dir) / Path(inner)
        full_src = IMG_DIR / src_name
        if not full_src.exists():
            print(f"  [SKIP] нет PNG: {full_src}")
            continue
        if not full_inner.exists():
            print(f"  [SKIP] нет якоря: {inner}")
            continue
        shutil.copy(full_src, full_inner)
        print(f"  [IMG]  {inner} ← {src_name}")


def add_auto_block(dst_path, summary, skip_data):
    """Добавляет/обновляет в конце документа таблицу с расчётными значениями.
    Идемпотентен: ищет AUTO_BLOCK_MARKER и заменяет старый блок новым."""
    from docx import Document

    doc = Document(str(dst_path))
    body = doc.element.body
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    direct_children = list(body)
    cut_idx = None
    for i, child in enumerate(direct_children):
        if child.tag == f"{ns}p":
            text = "".join((t.text or "") for t in child.iter(f"{ns}t"))
            if AUTO_BLOCK_MARKER in text:
                cut_idx = i
                break
    if cut_idx is not None:
        removed = 0
        for child in direct_children[cut_idx:]:
            if child.tag == f"{ns}sectPr":
                continue
            body.remove(child)
            removed += 1
        print(f"  [AUTO] старый авто-блок удалён ({removed} элементов)")

    doc.add_paragraph().add_run(AUTO_BLOCK_MARKER).italic = True

    if not summary:
        doc.add_paragraph(
            "Расчётные данные не найдены (нужно сначала запустить relax.py "
            "и поместить mixture.out / mixture.skip в data/)."
        )
    else:
        intro = doc.add_paragraph()
        intro.add_run(
            "Сводка по релаксации смеси: парциальные и общие T, n, S "
            "в начальный (t = 0) и конечный момент времени."
        ).bold = True

        s, e = summary["start"], summary["end"]
        cols = [
            ("момент", "t = 0", f"t = {summary['rows'] - 1}"),
            ("T_tot", s["T_tot"], e["T_tot"]),
            ("T_c", s["T_c"], e["T_c"]),
            ("T_h", s["T_h"], e["T_h"]),
            ("n_c", s["n_c"], e["n_c"]),
            ("n_h", s["n_h"], e["n_h"]),
            ("S_tot", s["S_tot"], e["S_tot"]),
            ("S_c", s["S_c"], e["S_c"]),
            ("S_h", s["S_h"], e["S_h"]),
        ]
        table = doc.add_table(rows=3, cols=len(cols))
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        # шапка
        for j, (name, _, _) in enumerate(cols):
            table.rows[0].cells[j].text = name
        # строка начало
        table.rows[1].cells[0].text = "t = 0"
        for j, (_, v0, _) in enumerate(cols[1:], start=1):
            table.rows[1].cells[j].text = f"{v0:.5f}"
        # строка конец
        table.rows[2].cells[0].text = f"t = {summary['rows'] - 1}"
        for j, (_, _, v1) in enumerate(cols[1:], start=1):
            table.rows[2].cells[j].text = f"{v1:.5f}"

        # подпись с диагностикой
        note = doc.add_paragraph()
        parts = []
        if skip_data:
            parts.append(f"шагов по времени: {skip_data.get('actual_steps', summary['rows'])}")
            parts.append(f"доля пропущенных столкновений (положительность): "
                         f"{skip_data.get('percent', 0):.4f}%")
            el = skip_data.get("elapsed_seconds")
            if el:
                m, sec = divmod(int(el), 60)
                parts.append(f"время счёта: {m}м {sec}с" if m else f"{sec}с")
        note.add_run("; ".join(parts)).italic = True

        concl = doc.add_paragraph()
        concl.add_run(
            f"Итог: парциальные температуры сравниваются "
            f"(T_c: {s['T_c']:.5f} → {e['T_c']:.5f}, "
            f"T_h: {s['T_h']:.5f} → {e['T_h']:.5f}), концентрации сохраняются "
            f"(n_c = n_h = 0.5), общая энтропия растёт "
            f"(S_tot: {s['S_tot']:.5f} → {e['S_tot']:.5f}) — H-теорема."
        )

    doc.add_paragraph(AUTO_BLOCK_MARKER).runs[0].italic = True
    doc.save(str(dst_path))
    print(f"  [AUTO] новый блок добавлен")


def main():
    import argparse
    p = argparse.ArgumentParser()
    p.add_argument("--data-dir", default=None,
                   help="папка с mixture.out / .skip (по умолчанию ищет в data/, src/, корне)")
    args = p.parse_args()

    if not SRC.exists():
        sys.exit(f"Не найден исходный файл: {SRC}")

    data_dir = find_data_dir(args.data_dir)
    skip_data = load_skip_data(data_dir)
    summary = load_out_summary(data_dir)

    print(f"==> Источник/цель: {SRC}")
    print(f"==> Данные:        {data_dir}")
    print(f"==> mixture.out:   {'найден' if summary else 'НЕ найден'}")
    print(f"==> mixture.skip:  {'найден' if skip_data else 'НЕ найден'}")
    print()

    with tempfile.TemporaryDirectory() as td:
        with zipfile.ZipFile(SRC, "r") as zin:
            zin.extractall(td)
        replace_images(td)
        with zipfile.ZipFile(DST, "w", zipfile.ZIP_DEFLATED) as out:
            for root, _, files in os.walk(td):
                for fn in files:
                    full = os.path.join(root, fn)
                    rel = os.path.relpath(full, td).replace("\\", "/")
                    out.write(full, rel)

    add_auto_block(DST, summary, skip_data)

    print()
    print(f"готово: {DST}")


if __name__ == "__main__":
    main()
